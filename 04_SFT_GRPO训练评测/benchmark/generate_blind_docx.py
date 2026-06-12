import json
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image
except Exception:
    Image = None


ROOT = Path.cwd()
BENCH_DIR = ROOT / "benchmark"
SOURCE_JSON = BENCH_DIR / "doctor_sample_20260610_blind.json"
OUT_DOCX = BENCH_DIR / "doctor_sample_20260610_blind.docx"


def set_east_asia_font(run, font_name):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def add_text(paragraph, text, *, bold=False, size=11, color=None):
    run = paragraph.add_run(text)
    set_east_asia_font(run, "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=14 if level == 1 else 10, after=7 if level == 1 else 5)
    run = add_text(paragraph, text, bold=True, size=16 if level == 1 else 13, color="2E74B5" if level == 1 else "1F4D78")
    return paragraph


def build_overview_table(doc, items):
    counts = {}
    for item in items:
        counts[item["modality"]] = counts.get(item["modality"], 0) + 1
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table, 9360)
    headers = ["题型", "题数"]
    widths = [4680, 4680]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_margins(cell)
        set_cell_shading(cell, "E8EEF5")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        add_text(p, header, bold=True)
    for modality in sorted(counts):
        row = table.add_row()
        values = [modality, str(counts[modality])]
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            if idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text(p, value)
    return table


def image_width(image_path):
    if not Image or not image_path or not os.path.exists(image_path):
        return Inches(3.0)
    try:
        with Image.open(image_path) as img:
            w, h = img.size
        if w <= 0 or h <= 0:
            return Inches(3.0)
        width_in = min(3.2, max(1.8, w / 180.0))
        if h / w > 1.15:
            width_in = min(width_in, 2.5)
        return Inches(width_in)
    except Exception:
        return Inches(3.0)


def add_question(doc, item):
    q_head = doc.add_paragraph()
    q_head.paragraph_format.keep_with_next = True
    set_paragraph_spacing(q_head, before=10, after=4, line=1.15)
    add_text(q_head, f"第 {item['sample_no']} 题", bold=True, size=12, color="7B3F2C")
    add_text(q_head, f"    {item['modality']}", size=10, color="636B60")

    image_path = item.get("image_abs_path") or ""
    if image_path and os.path.exists(image_path):
        p_img = doc.add_paragraph()
        p_img.paragraph_format.keep_with_next = True
        p_img.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p_img, before=2, after=4)
        run = p_img.add_run()
        run.add_picture(image_path, width=image_width(image_path))

    for idx, line in enumerate(str(item["question"]).splitlines()):
        p = doc.add_paragraph()
        if idx == 0:
            p.paragraph_format.keep_with_next = True
        set_paragraph_spacing(p, before=0, after=4)
        add_text(p, line, bold=idx == 0)

    for letter in ["A", "B", "C", "D"]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        set_paragraph_spacing(p, before=0, after=3)
        add_text(p, f"{letter}. ", bold=True, color="2F6756")
        add_text(p, item["options"].get(letter, ""))

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=3, after=8)
    add_text(p, "作答：□ A　□ B　□ C　□ D        题目问题：□ 无　□ 有", size=10, color="636B60")


def build_docx():
    data = json.loads(SOURCE_JSON.read_text(encoding="utf-8"))
    items = data["items"]

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=0, after=12)
    add_text(title, "中医 Benchmark 医生盲测题本", bold=True, size=20, color="1F4D78")

    add_heading(doc, "题本概览", level=1)
    build_overview_table(doc, items)
    add_heading(doc, "正式题目", level=1)

    for item in items:
        add_question(doc, item)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(footer, "中医 Benchmark 医生盲测题本", size=9, color="636B60")

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_docx()
